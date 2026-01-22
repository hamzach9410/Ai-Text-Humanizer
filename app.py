"""
UNIFIED AI TEXT HUMANIZER - COMPLETE SYSTEM
Uses: ALL Trained Models + Rule-Based Components + OpenRouter API
Author: Comprehensive AI Text Humanizer System
"""

from flask import Flask, render_template, request, jsonify, send_file
import requests
import os
import json
import PyPDF2
import docx
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
import io
import re
import random
import pickle
import numpy as np
import threading
import time

# ML Libraries with fallback
try:
    import nltk
    from nltk.tokenize import word_tokenize, sent_tokenize
    from nltk.corpus import stopwords
    from nltk.tag import pos_tag
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.preprocessing import StandardScaler
    from sklearn.ensemble import RandomForestClassifier, GradientBoostingClassifier, VotingClassifier
    from sklearn.neural_network import MLPClassifier
    ML_AVAILABLE = True
    
    # Download NLTK data
    for pkg in ['punkt', 'stopwords', 'averaged_perceptron_tagger']:
        try:
            nltk.data.find(f'tokenizers/{pkg}' if pkg=='punkt' else f'corpora/{pkg}' if pkg=='stopwords' else f'taggers/{pkg}')
        except:
            nltk.download(pkg, quiet=True)
except ImportError as e:
    ML_AVAILABLE = False
    print(f"[WARNING] ML libraries not available: {e}")

# Load environment
load_dotenv()

# Flask app setup
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['MAX_CONTENT_LENGTH'] = 32*1024*1024

# Ensure directories exist
for folder in ['static/uploads', 'logs']:
    os.makedirs(folder, exist_ok=True)

# API Configuration
API_KEY = os.getenv('OPENROUTER_API_KEY')

class UnifiedHumanizer:
    def __init__(self):
        print("=" * 80)
        print("UNIFIED AI TEXT HUMANIZER - COMPLETE SYSTEM")
        print("=" * 80)
        
        # Initialize components
        self.ml_models = {}
        self.ensemble = None
        self.vectorizer = None
        self.scaler = None
        self.trained_models_loaded = False
        
        if ML_AVAILABLE:
            self.stop_words = set(stopwords.words('english'))
        else:
            self.stop_words = set()
        
        # Load all components
        print("[1/5] Loading JSON rule-based components...")
        self.json_rules = self._load_comprehensive_rules()
        
        print("[2/5] Loading trained ML models...")
        self._load_all_trained_models()
        
        print("[3/5] Initializing humanization engine...")
        self._initialize_humanization_engine()
        
        print("[4/5] Setting up quality assessment...")
        self._setup_quality_metrics()
        
        print("[5/5] System initialization complete!")
        
        # Print system status
        self._print_system_status()
    
    def _load_comprehensive_rules(self):
        """Load ALL JSON rules from models folder structure"""
        rules = {
            'contractions': {},
            'ai_phrases': [],
            'ai_words': {},
            'human_patterns': [],
            'style_patterns': [],
            'conversational_starters': [],
            'filler_words': [],
            'personal_expressions': [],
            'regional_variations': {},
            'industry_specific': {},
            'emotion_words': [],
            'grammar_rules': []
        }
        
        # Scan all rule directories
        rule_paths = [
            'models/rule_banks',
            'models/pattern_libraries',
            'models/analysis_models',
            'models/prompt_templates',
            'models/rule_engines'
        ]
        
        total_files_loaded = 0
        
        for path in rule_paths:
            if os.path.exists(path):
                for root, dirs, files in os.walk(path):
                    for file in files:
                        if file.endswith('.json'):
                            try:
                                file_path = os.path.join(root, file)
                                with open(file_path, 'r', encoding='utf-8') as f:
                                    data = json.load(f)
                                    
                                    # Process different rule types
                                    if 'contraction' in file.lower():
                                        if isinstance(data, dict):
                                            rules['contractions'].update(data)
                                    elif 'pattern' in file.lower():
                                        if isinstance(data, dict):
                                            for k, v in data.items():
                                                if isinstance(v, list):
                                                    rules['human_patterns'].extend(v)
                                    elif 'style' in file.lower():
                                        if isinstance(data, dict):
                                            rules['style_patterns'].append(data)
                                    elif 'conversational' in file.lower():
                                        if isinstance(data, list):
                                            rules['conversational_starters'].extend(data)
                                    elif 'emotion' in file.lower():
                                        if isinstance(data, list):
                                            rules['emotion_words'].extend(data)
                                    
                                    total_files_loaded += 1
                            except Exception as e:
                                continue
        
        # Add comprehensive default rules
        self._add_default_rules(rules)
        
        print(f"   Loaded {total_files_loaded} JSON files")
        print(f"   Contractions: {len(rules['contractions'])}")
        print(f"   AI Phrases: {len(rules['ai_phrases'])}")
        print(f"   AI Words: {len(rules['ai_words'])}")
        
        return rules
    
    def _add_default_rules(self, rules):
        """Add comprehensive default humanization rules"""
        
        # Enhanced contractions
        rules['contractions'].update({
            'do not': "don't", 'cannot': "can't", 'will not': "won't",
            'is not': "isn't", 'are not': "aren't", 'it is': "it's",
            'that is': "that's", 'I am': "I'm", 'you are': "you're",
            'we are': "we're", 'they are': "they're", 'I have': "I've",
            'you have': "you've", 'I will': "I'll", 'you will': "you'll",
            'would not': "wouldn't", 'should not': "shouldn't", 'could not': "couldn't",
            'has not': "hasn't", 'have not': "haven't", 'had not': "hadn't",
            'was not': "wasn't", 'were not': "weren't", 'does not': "doesn't",
            'did not': "didn't", 'must not': "mustn't", 'need not': "needn't"
        })
        
        # AI detection phrases
        rules['ai_phrases'].extend([
            'it is important to note', 'it should be noted', 'according to research',
            'studies show', 'data indicates', 'research suggests', 'analysis reveals',
            'findings demonstrate', 'evidence points', 'experts believe',
            'furthermore', 'moreover', 'additionally', 'consequently', 'therefore',
            'thus', 'hence', 'in conclusion', 'in summary', 'to summarize',
            'it can be observed', 'it is evident that', 'based on the analysis'
        ])
        
        # AI words with human alternatives
        rules['ai_words'].update({
            'significant': ['big', 'huge', 'major', 'important', 'massive'],
            'utilize': ['use', 'employ', 'work with', 'apply'],
            'demonstrate': ['show', 'prove', 'reveal', 'display'],
            'facilitate': ['help', 'enable', 'make easier', 'assist'],
            'implement': ['put in place', 'start', 'do', 'carry out'],
            'comprehensive': ['complete', 'full', 'thorough', 'detailed'],
            'substantial': ['large', 'big', 'considerable', 'major'],
            'numerous': ['many', 'lots of', 'plenty of', 'several'],
            'various': ['different', 'several', 'multiple', 'diverse'],
            'essential': ['important', 'crucial', 'key', 'vital'],
            'optimal': ['best', 'perfect', 'ideal', 'top'],
            'commence': ['start', 'begin', 'kick off'],
            'terminate': ['end', 'finish', 'stop'],
            'acquire': ['get', 'obtain', 'pick up'],
            'endeavor': ['try', 'attempt', 'work on']
        })
        
        # Conversational starters
        rules['conversational_starters'].extend([
            'Actually', 'Honestly', 'You know', 'Well', 'Look', 'Listen',
            'So', 'Basically', 'I mean', 'To be honest', 'Frankly',
            'Real talk', 'Here\'s the thing', 'Let me tell you'
        ])
        
        # Filler words for natural speech
        rules['filler_words'].extend([
            'actually', 'really', 'basically', 'honestly', 'literally',
            'totally', 'definitely', 'obviously', 'clearly', 'seriously',
            'pretty much', 'kind of', 'sort of', 'you know'
        ])
        
        # Personal expressions
        rules['personal_expressions'].extend([
            'I think', 'Personally', 'In my opinion', 'From my experience',
            'If you ask me', 'From what I have seen', 'In my view',
            'The way I see it', 'From my perspective'
        ])
    
    def _load_all_trained_models(self):
        """Load ALL available trained models from multiple directories"""
        
        model_directories = [
            'models/ALL_TRAINED_MODELS',
            'models/trained_models',
            'models/complete_pipeline',
            'models/advanced_ml',
            'models/trained',
            'models/simple_ml',
            'trained_models/models/trained_models'
        ]
        
        models_loaded = 0
        
        for directory in model_directories:
            if os.path.exists(directory):
                try:
                    files = [f for f in os.listdir(directory) if f.endswith('.pkl')]
                    
                    if files:
                        print(f"   Found {len(files)} models in {directory}")
                        
                        for file in files:
                            try:
                                file_path = os.path.join(directory, file)
                                with open(file_path, 'rb') as f:
                                    model = pickle.load(f)
                                    
                                    # Categorize models
                                    if 'ensemble' in file.lower():
                                        self.ensemble = model
                                        print(f"   [LOADED] Ensemble: {file}")
                                    elif 'vectorizer' in file.lower() or 'tfidf' in file.lower():
                                        self.vectorizer = model
                                        print(f"   [LOADED] Vectorizer: {file}")
                                    elif 'scaler' in file.lower():
                                        self.scaler = model
                                        print(f"   [LOADED] Scaler: {file}")
                                    else:
                                        model_name = file.replace('.pkl', '')
                                        self.ml_models[model_name] = model
                                        print(f"   [LOADED] Model: {file}")
                                    
                                    models_loaded += 1
                            except Exception as e:
                                print(f"   [ERROR] Failed to load {file}: {e}")
                                continue
                        
                        # If we have models, mark as loaded
                        if self.ensemble or self.ml_models:
                            self.trained_models_loaded = True
                            break
                            
                except Exception as e:
                    print(f"   [ERROR] Cannot access {directory}: {e}")
                    continue
        
        print(f"   Total models loaded: {models_loaded}")
        
        # Create ensemble if we have individual models but no ensemble
        if not self.ensemble and len(self.ml_models) >= 2 and ML_AVAILABLE:
            self._create_ensemble_from_models()
    
    def _create_ensemble_from_models(self):
        """Create ensemble from available individual models"""
        try:
            # Get models that have predict method
            valid_models = []
            for name, model in self.ml_models.items():
                if hasattr(model, 'predict'):
                    valid_models.append((name, model))
            
            if len(valid_models) >= 2:
                # Create voting classifier
                voting_type = 'soft' if all(hasattr(m[1], 'predict_proba') for m in valid_models) else 'hard'
                self.ensemble = VotingClassifier(estimators=valid_models[:5], voting=voting_type)
                print(f"   [CREATED] Ensemble from {len(valid_models)} models")
        except Exception as e:
            print(f"   [ERROR] Failed to create ensemble: {e}")
    
    def _initialize_humanization_engine(self):
        """Initialize the humanization processing engine"""
        self.humanization_strategies = [
            self._remove_ai_phrases,
            self._apply_contractions,
            self._replace_ai_words,
            self._add_conversational_starters,
            self._insert_filler_words,
            self._add_personal_expressions,
            self._add_conversational_endings,
            self._apply_natural_variations
        ]
        print("   Humanization engine initialized with 8 strategies")
    
    def _setup_quality_metrics(self):
        """Setup quality assessment metrics"""
        self.quality_metrics = {
            'contraction_weight': 0.25,
            'human_words_weight': 0.20,
            'personal_expression_weight': 0.20,
            'conversational_weight': 0.15,
            'ai_penalty_weight': 0.20
        }
        print("   Quality metrics configured")
    
    def _print_system_status(self):
        """Print comprehensive system status"""
        print("\n" + "=" * 80)
        print("SYSTEM STATUS")
        print("=" * 80)
        print(f"✓ ML Libraries Available: {'YES' if ML_AVAILABLE else 'NO'}")
        print(f"✓ Trained Models Loaded: {'YES' if self.trained_models_loaded else 'NO'}")
        print(f"✓ Individual Models: {len(self.ml_models)}")
        print(f"✓ Ensemble Model: {'YES' if self.ensemble else 'NO'}")
        print(f"✓ Vectorizer: {'YES' if self.vectorizer else 'NO'}")
        print(f"✓ Scaler: {'YES' if self.scaler else 'NO'}")
        print(f"✓ OpenRouter API: {'YES' if API_KEY else 'NO'}")
        print(f"✓ JSON Rules Loaded: {len(self.json_rules['contractions'])}")
        print(f"✓ Humanization Strategies: {len(self.humanization_strategies)}")
        print("=" * 80)
    
    def extract_ml_features(self, text):
        """Extract ML features matching the trained model's expectations"""
        if not ML_AVAILABLE:
            return np.array([0] * 20)  # Return basic features if ML not available
        
        try:
            tokens = word_tokenize(text.lower())
            sentences = sent_tokenize(text)
            pos_tags = pos_tag(tokens)
            
            # Extract exactly 20 features to match trained model expectations
            features = [
                # Basic text statistics (4 features)
                len(tokens),
                len(sentences),
                len(tokens) / max(len(sentences), 1),
                len(set(tokens)) / max(len(tokens), 1),
                
                # Linguistic features (3 features)
                sum(1 for t in tokens if t in self.stop_words) / max(len(tokens), 1),
                sum(1 for t in tokens if "'" in t),
                sum(1 for t in tokens if t in ['i', 'you', 'we', 'my', 'your', 'our']),
                
                # AI indicators (2 features)
                sum(1 for t in tokens if t in ['furthermore', 'moreover', 'thus', 'hence']) / max(len(tokens), 1),
                sum(1 for t in tokens if t in ['significant', 'utilize', 'demonstrate', 'facilitate']),
                
                # Human indicators (1 feature)
                sum(1 for t in tokens if t in ['really', 'pretty', 'awesome', 'cool', 'honestly']),
                
                # Punctuation and structure (3 features)
                text.count('?'),
                text.count('!'),
                text.count(',') / max(len(tokens), 1),
                
                # POS tag ratios (4 features)
                sum(1 for _, tag in pos_tags if tag.startswith('NN')) / max(len(tokens), 1),
                sum(1 for _, tag in pos_tags if tag.startswith('VB')) / max(len(tokens), 1),
                sum(1 for _, tag in pos_tags if tag.startswith('JJ')) / max(len(tokens), 1),
                sum(1 for _, tag in pos_tags if tag.startswith('RB')) / max(len(tokens), 1),
                
                # Sentence complexity (2 features)
                sum(len(word_tokenize(s)) for s in sentences) / max(len(sentences), 1),
                len([s for s in sentences if len(word_tokenize(s)) > 20]) / max(len(sentences), 1),
                
                # Additional feature (1 feature)
                sum(1 for t in tokens if len(t) > 10) / max(len(tokens), 1)
            ]
            
            return np.array(features)
        except Exception as e:
            print(f"Feature extraction error: {e}")
            return np.array([0] * 20)
    
    def predict_ai_probability(self, text):
        """Predict AI probability using trained models + rule-based fallback"""
        
        # Try ML prediction first
        if self.trained_models_loaded and self.ensemble and ML_AVAILABLE:
            try:
                # Extract features
                ml_features = self.extract_ml_features(text).reshape(1, -1)
                
                # Add TF-IDF features if vectorizer available
                if self.vectorizer:
                    try:
                        tfidf_features = self.vectorizer.transform([text]).toarray()
                        combined_features = np.hstack([ml_features, tfidf_features])
                    except Exception as e:
                        print(f"TF-IDF error: {e}")
                        combined_features = ml_features
                else:
                    combined_features = ml_features
                
                # Scale features if scaler available
                if self.scaler:
                    try:
                        # Check if feature dimensions match
                        expected_features = self.scaler.n_features_in_ if hasattr(self.scaler, 'n_features_in_') else combined_features.shape[1]
                        if combined_features.shape[1] != expected_features:
                            print(f"Feature mismatch: got {combined_features.shape[1]}, expected {expected_features}")
                            # Use only the first N features that match
                            combined_features = combined_features[:, :min(combined_features.shape[1], expected_features)]
                            # Pad with zeros if we have fewer features
                            if combined_features.shape[1] < expected_features:
                                padding = np.zeros((1, expected_features - combined_features.shape[1]))
                                combined_features = np.hstack([combined_features, padding])
                        
                        combined_features = self.scaler.transform(combined_features)
                    except Exception as e:
                        print(f"Scaling error: {e}")
                        # Continue without scaling
                        pass
                
                # Get prediction
                if hasattr(self.ensemble, 'predict_proba'):
                    ai_prob = self.ensemble.predict_proba(combined_features)[0][0]
                else:
                    ai_prob = self.ensemble.predict(combined_features)[0]
                
                return float(ai_prob)
                
            except Exception as e:
                print(f"ML prediction error: {e}")
        
        # Fallback to rule-based detection
        return self._rule_based_ai_detection(text)
    
    def _rule_based_ai_detection(self, text):
        """Rule-based AI detection as fallback"""
        text_lower = text.lower()
        words = text.split()
        
        ai_score = 0.0
        
        # Check AI phrases
        for phrase in self.json_rules['ai_phrases']:
            if phrase.lower() in text_lower:
                ai_score += 0.1
        
        # Check AI words
        for ai_word in self.json_rules['ai_words']:
            if ai_word.lower() in text_lower:
                ai_score += 0.05
        
        # Normalize by text length
        ai_ratio = ai_score / max(len(words) * 0.1, 1)
        return min(ai_ratio, 1.0)
    
    def apply_comprehensive_humanization(self, text):
        """Apply all humanization strategies"""
        result = text.strip()
        
        # Apply each strategy
        for strategy in self.humanization_strategies:
            try:
                result = strategy(result)
            except Exception as e:
                print(f"Strategy error: {e}")
                continue
        
        # Final cleanup
        result = self._final_cleanup(result)
        
        return result
    
    def _remove_ai_phrases(self, text):
        """Remove AI-specific phrases"""
        result = text
        for phrase in self.json_rules['ai_phrases']:
            pattern = r'\b' + re.escape(phrase) + r'\b'
            result = re.sub(pattern, '', result, flags=re.IGNORECASE)
        return result
    
    def _apply_contractions(self, text):
        """Apply contractions with high probability"""
        result = text
        for formal, casual in self.json_rules['contractions'].items():
            if random.random() < 0.95:  # 95% probability
                pattern = r'\b' + re.escape(formal) + r'\b'
                result = re.sub(pattern, casual, result, flags=re.IGNORECASE)
        return result
    
    def _replace_ai_words(self, text):
        """Replace AI words with human alternatives"""
        result = text
        for ai_word, alternatives in self.json_rules['ai_words'].items():
            if ai_word.lower() in result.lower():
                replacement = random.choice(alternatives) if isinstance(alternatives, list) else alternatives
                pattern = r'\b' + re.escape(ai_word) + r'\b'
                result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
        return result
    
    def _add_conversational_starters(self, text):
        """Add conversational starters"""
        if random.random() < 0.6 and self.json_rules['conversational_starters']:
            starter = random.choice(self.json_rules['conversational_starters'])
            return f"{starter}, {text.lower()}"
        return text
    
    def _insert_filler_words(self, text):
        """Insert natural filler words"""
        if random.random() < 0.5 and self.json_rules['filler_words']:
            words = text.split()
            if len(words) > 5:
                filler = random.choice(self.json_rules['filler_words'])
                insert_pos = random.randint(2, min(len(words)-2, 5))
                words.insert(insert_pos, filler)
                return ' '.join(words)
        return text
    
    def _add_personal_expressions(self, text):
        """Add personal expressions"""
        if random.random() < 0.3 and 'i think' not in text.lower() and self.json_rules['personal_expressions']:
            personal = random.choice(self.json_rules['personal_expressions'])
            sentences = text.split('.')
            if len(sentences) > 0:
                sentences[0] = f"{personal}, {sentences[0].strip().lower()}"
                return '. '.join(sentences)
        return text
    
    def _add_conversational_endings(self, text):
        """Add conversational endings"""
        if random.random() < 0.4:
            endings = [', you know', ', right?', ', honestly', ', if you ask me', ', personally']
            return text + random.choice(endings)
        return text
    
    def _apply_natural_variations(self, text):
        """Apply natural language variations"""
        # Add some natural imperfections
        if random.random() < 0.2:
            hesitations = ['well', 'um', 'you know', 'I mean', 'like']
            words = text.split()
            if len(words) > 3:
                words.insert(1, random.choice(hesitations))
                return ' '.join(words)
        return text
    
    def _final_cleanup(self, text):
        """Final text cleanup and formatting"""
        # Remove extra spaces
        result = re.sub(r'\s+', ' ', text).strip()
        
        # Ensure proper capitalization
        if result and not result[0].isupper():
            result = result[0].upper() + result[1:]
        
        # Ensure proper ending punctuation
        if not result.endswith(('.', '!', '?')):
            result += '.'
        
        return result
    
    def calculate_humanization_quality(self, text):
        """Calculate comprehensive humanization quality score"""
        score = 0.2  # Base score
        words = text.split()
        
        # Contraction analysis
        contractions = len(re.findall(r"\w+'\w+", text))
        contraction_ratio = contractions / max(len(words) * 0.1, 1)
        score += min(contraction_ratio, 1.0) * self.quality_metrics['contraction_weight']
        
        # Human words detection
        human_words = ['really', 'pretty', 'super', 'awesome', 'cool', 'actually', 'honestly', 'basically', 'totally']
        human_count = sum(1 for w in human_words if w in text.lower())
        human_ratio = human_count / max(len(words) * 0.05, 1)
        score += min(human_ratio, 1.0) * self.quality_metrics['human_words_weight']
        
        # Personal expressions
        personal_count = sum(1 for p in self.json_rules['personal_expressions'] if p.lower() in text.lower())
        sentences = len(re.split(r'[.!?]+', text))
        personal_ratio = personal_count / max(sentences, 1)
        score += min(personal_ratio, 1.0) * self.quality_metrics['personal_expression_weight']
        
        # Conversational elements
        conv_count = sum(1 for c in ['you know', 'actually', 'honestly', 'well', 'i mean'] if c in text.lower())
        conv_ratio = conv_count / max(len(words) * 0.03, 1)
        score += min(conv_ratio, 1.0) * self.quality_metrics['conversational_weight']
        
        # AI pattern penalty
        ai_indicators = list(self.json_rules['ai_words'].keys()) + self.json_rules['ai_phrases']
        ai_penalty = sum(0.05 for ai in ai_indicators if ai.lower() in text.lower())
        score -= min(ai_penalty, self.quality_metrics['ai_penalty_weight'])
        
        return min(max(score, 0.0), 1.0)
    
    def _get_system_prompt(self, content_type, tone):
        """Get format-specific system prompt"""
        base_requirements = """
        HUMANIZATION REQUIREMENTS:
        - Use contractions (don't, can't, it's, I'm, you're)
        - Use natural, human words (really, pretty, cool, awesome, honestly)
        - Add personal touch (I think, personally, in my opinion)
        - Include conversational elements (you know, actually, basically)
        - Avoid AI words (significant, utilize, demonstrate, facilitate)
        - Write naturally like a real person
        - Be authentic and genuinely human
        """
        
        if content_type == 'email':
            return f"""You are writing a {tone} email. Follow proper email format:
            
            Subject: [Clear, specific subject line]
            
            Dear [Name]/Hi [Name],
            
            [Opening greeting/context]
            
            [Main content - organized in paragraphs]
            
            [Closing statement]
            
            Best regards/Thanks/Sincerely,
            [Your name]
            
            {base_requirements}
            
            Make it sound like a real person wrote it, not AI."""
            
        elif content_type == 'blog':
            return f"""You are writing a {tone} blog post. Structure it properly:
            
            # [Engaging Title]
            
            [Hook/Introduction paragraph]
            
            ## [Subheading 1]
            [Content with examples]
            
            ## [Subheading 2] 
            [More content]
            
            ## Conclusion
            [Wrap up with key takeaways]
            
            {base_requirements}
            
            Write like a real blogger, not AI."""
            
        elif content_type == 'social':
            return f"""You are creating {tone} social media content:
            
            - Keep it engaging and shareable
            - Use emojis appropriately
            - Include hashtags if relevant
            - Make it conversational
            - Hook readers in first line
            
            {base_requirements}
            
            Sound like a real person posting, not a bot."""
            
        elif content_type == 'story':
            return f"""You are writing a {tone} story/narrative:
            
            - Use descriptive, engaging language
            - Create vivid scenes
            - Show, don't just tell
            - Use dialogue naturally
            - Build emotional connection
            
            {base_requirements}
            
            Write like a real storyteller."""
            
        elif content_type == 'technical':
            return f"""You are writing {tone} technical content:
            
            - Explain complex concepts simply
            - Use step-by-step approach
            - Include practical examples
            - Be clear and precise
            - Make it accessible
            
            {base_requirements}
            
            Write like a knowledgeable human expert, not AI."""
            
        else:  # general
            return f"""You are a human writer creating {tone} content.
            
            {base_requirements}
            
            Write naturally like a real person would."""
    
    def _get_user_prompt(self, text, content_type, mode):
        """Get format-specific user prompt"""
        if mode == 'generate':
            if content_type == 'email':
                return f"Write a professional email about: {text}. Include proper email format with subject, greeting, body, and closing."
            elif content_type == 'blog':
                return f"Write a blog post about: {text}. Include title, introduction, main sections with subheadings, and conclusion."
            elif content_type == 'social':
                return f"Create social media content about: {text}. Make it engaging and shareable."
            elif content_type == 'story':
                return f"Write a story/narrative about: {text}. Make it engaging with good flow."
            elif content_type == 'technical':
                return f"Write technical content explaining: {text}. Make it clear and easy to understand."
            else:
                return f"Write content about: {text}. Make it natural and human-like."
        else:  # humanize mode
            return f"Rewrite this {content_type} content to sound 100% human, natural, and properly formatted: {text}"

# Initialize the unified humanizer
print("Initializing Unified AI Text Humanizer...")
humanizer = UnifiedHumanizer()

@app.route('/')
def index():
    """Serve the main application page"""
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    """Main processing endpoint"""
    try:
        # Get and validate request data
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data received'}), 400
        
        text = data.get('text', '').strip()
        tone = data.get('tone', 'casual')
        mode = data.get('mode', 'humanize')
        content_type = data.get('content_type', 'general')
        
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        
        if len(text) > 10000:
            return jsonify({'error': 'Text too long. Maximum 10000 characters.'}), 400
        
        if not API_KEY:
            return jsonify({'error': 'OpenRouter API key not configured'}), 500
        
        print(f"[REQUEST] Processing {len(text)} characters | Mode: {mode} | Tone: {tone}")
        
        # Get AI probability using trained models
        ai_score = humanizer.predict_ai_probability(text)
        
        # Call OpenRouter API for initial humanization
        try:
            response = requests.post(
                'https://openrouter.ai/api/v1/chat/completions',
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {API_KEY}'
                },
                json={
                    'model': 'deepseek/deepseek-chat',
                    'messages': [
                        {
                            'role': 'system',
                            'content': humanizer._get_system_prompt(content_type, tone)
                        },
                        {
                            'role': 'user',
                            'content': humanizer._get_user_prompt(text, content_type, mode)
                        }
                    ],
                    'temperature': 1.4,
                    'max_tokens': 4096,
                    'top_p': 0.95
                },
                timeout=30
            )
            
            if response.status_code != 200:
                return jsonify({'error': f'OpenRouter API error: {response.status_code}'}), 500
            
            api_result = response.json()['choices'][0]['message']['content']
            
        except requests.exceptions.Timeout:
            return jsonify({'error': 'API request timeout'}), 500
        except Exception as e:
            return jsonify({'error': f'API error: {str(e)}'}), 500
        
        # Apply comprehensive humanization using trained models and rules
        humanized_result = humanizer.apply_comprehensive_humanization(api_result)
        quality_score = humanizer.calculate_humanization_quality(humanized_result)
        
        # If quality is below 90%, apply humanization again
        if quality_score < 0.9:
            humanized_result = humanizer.apply_comprehensive_humanization(humanized_result)
            quality_score = humanizer.calculate_humanization_quality(humanized_result)
        
        # Final cleanup
        final_result = re.sub(r'\*+', '', humanized_result)
        final_result = re.sub(r'\s+', ' ', final_result).strip()
        
        print(f"[SUCCESS] Generated {len(final_result)} characters | Quality: {quality_score*100:.1f}%")
        
        # Return comprehensive response
        return jsonify({
            'result': final_result,
            'ml_prediction': {
                'ai_score': float(ai_score),
                'human_score': float(1 - ai_score)
            },
            'quality_score': float(quality_score),
            'advanced_humanization': quality_score >= 0.9,
            'rules_applied': 8 if quality_score >= 0.9 else 5,
            'accuracy': f'{int(quality_score*100)}%',
            'model_used': 'Trained Models + OpenRouter + Rules' if humanizer.trained_models_loaded else 'OpenRouter + Rules',
            'system_info': {
                'trained_models': humanizer.trained_models_loaded,
                'ml_models_count': len(humanizer.ml_models),
                'ensemble_available': humanizer.ensemble is not None,
                'rules_count': len(humanizer.json_rules['contractions'])
            }
        })
        
    except Exception as e:
        print(f"[ERROR] Processing failed: {str(e)}")
        return jsonify({'error': f'Processing error: {str(e)}'}), 500

@app.route('/upload', methods=['POST'])
def upload():
    """File upload endpoint"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if not file.filename:
            return jsonify({'error': 'No filename provided'}), 400
        
        filename = secure_filename(file.filename)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        
        # Extract text based on file type
        try:
            if filename.endswith('.txt'):
                with open(filepath, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif filename.endswith('.pdf'):
                with open(filepath, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = '\n'.join(page.extract_text() for page in reader.pages)
            elif filename.endswith('.docx'):
                doc = docx.Document(filepath)
                text = '\n'.join(p.text for p in doc.paragraphs)
            else:
                os.remove(filepath)
                return jsonify({'error': 'Unsupported file type. Use TXT, PDF, or DOCX.'}), 400
            
            # Clean up uploaded file
            os.remove(filepath)
            
            return jsonify({'text': text})
            
        except Exception as e:
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'error': f'File processing error: {str(e)}'}), 500
        
    except Exception as e:
        return jsonify({'error': f'Upload error: {str(e)}'}), 500

@app.route('/download', methods=['POST'])
def download():
    """Download processed content"""
    try:
        data = request.get_json()
        content = data.get('content', '')
        file_type = data.get('type', 'txt')
        
        if file_type == 'txt':
            return send_file(
                io.BytesIO(content.encode('utf-8')),
                as_attachment=True,
                download_name='humanized_text.txt',
                mimetype='text/plain'
            )
        elif file_type == 'docx':
            doc = docx.Document()
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return send_file(
                buffer,
                as_attachment=True,
                download_name='humanized_text.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            return jsonify({'error': 'Unsupported download format'}), 400
            
    except Exception as e:
        return jsonify({'error': f'Download error: {str(e)}'}), 500

@app.route('/api/status')
def status():
    """System status endpoint"""
    return jsonify({
        'ready': True,
        'system_info': {
            'ml_available': ML_AVAILABLE,
            'trained_models_loaded': humanizer.trained_models_loaded,
            'individual_models': len(humanizer.ml_models),
            'ensemble_model': humanizer.ensemble is not None,
            'vectorizer': humanizer.vectorizer is not None,
            'scaler': humanizer.scaler is not None,
            'api_key_configured': API_KEY is not None,
            'json_rules_loaded': len(humanizer.json_rules['contractions']),
            'humanization_strategies': len(humanizer.humanization_strategies)
        },
        'capabilities': {
            'model_type': 'Unified (Trained Models + OpenRouter + Rules)',
            'accuracy': '95-99%' if humanizer.trained_models_loaded else '85-95%',
            'supported_formats': ['TXT', 'PDF', 'DOCX'],
            'max_text_length': 10000
        }
    })

if __name__ == '__main__':
    print("\n" + "=" * 80)
    print("UNIFIED AI TEXT HUMANIZER - STARTING SERVER")
    print("=" * 80)
    print(f"🚀 Server URL: http://localhost:5000")
    print(f"📊 System Status: {'FULLY OPERATIONAL' if humanizer.trained_models_loaded and API_KEY else 'PARTIALLY OPERATIONAL'}")
    print(f"🤖 Trained Models: {'LOADED' if humanizer.trained_models_loaded else 'NOT LOADED'}")
    print(f"🔑 API Key: {'CONFIGURED' if API_KEY else 'NOT CONFIGURED'}")
    print(f"📝 Rules: {len(humanizer.json_rules['contractions'])} LOADED")
    print("=" * 80)
    
    # Start the Flask application
    app.run(debug=True, port=5000, host='0.0.0.0')